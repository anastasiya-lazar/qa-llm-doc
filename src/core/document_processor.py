from typing import List, Dict, Any
import uuid
from datetime import datetime
from pathlib import Path
import PyPDF2
from pydantic import BaseModel
from io import BytesIO
from docx import Document

from src.schemas.main_schemas import DocumentType, DocumentMetadata, DocumentChunk

class DocumentProcessor:
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def process_document(self, file_content: bytes, filename: str) -> DocumentMetadata:
        """Process an uploaded document and extract its content and metadata."""
        document_id = str(uuid.uuid4())
        file_path = self.upload_dir / f"{document_id}_{filename}"
        
        # Save the file
        with open(file_path, "wb") as f:
            f.write(file_content)

        # Determine document type
        doc_type = self._get_document_type(filename)
        
        # Extract text and metadata based on document type
        if doc_type == DocumentType.PDF:
            text, metadata = self._process_pdf(file_path)
        elif doc_type == DocumentType.TXT:
            text, metadata = self._process_txt(file_path)
        elif doc_type == DocumentType.DOCX:
            text, metadata = self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

        # Create document metadata
        document_metadata = DocumentMetadata(
            document_id=document_id,
            filename=filename,
            document_type=doc_type,
            page_count=metadata.get("page_count"),
            word_count=metadata.get("word_count"),
            chunk_count=0  # Will be updated after chunking
        )

        return document_metadata, text

    async def process_document_in_memory(self, file_content: bytes, filename: str) -> DocumentMetadata:
        """Process a document in memory without saving it to disk."""
        if not filename:
            raise ValueError("Filename cannot be empty")
            
        # Determine document type
        try:
            doc_type = self._get_document_type(filename)
        except ValueError as e:
            raise ValueError(f"Invalid document type: {str(e)}")
        
        # Extract text and metadata based on document type
        try:
            if doc_type == DocumentType.PDF:
                text, metadata = self._process_pdf_in_memory(file_content)
            elif doc_type == DocumentType.TXT:
                text, metadata = self._process_txt_in_memory(file_content)
            elif doc_type == DocumentType.DOCX:
                text, metadata = self._process_docx_in_memory(file_content)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
        except Exception as e:
            raise ValueError(f"Error processing document: {str(e)}")

        # Create document metadata
        document_metadata = DocumentMetadata(
            document_id=str(uuid.uuid4()),
            filename=filename,
            document_type=doc_type,
            page_count=metadata.get("page_count"),
            word_count=metadata.get("word_count"),
            chunk_count=0  # Will be updated after chunking
        )

        return document_metadata, text

    def _get_document_type(self, filename: str) -> DocumentType:
        """Determine the document type from the filename."""
        if not filename:
            raise ValueError("Filename cannot be empty")
            
        ext = filename.lower().split(".")[-1]
        if ext == "pdf":
            return DocumentType.PDF
        elif ext == "txt":
            return DocumentType.TXT
        elif ext == "docx":
            return DocumentType.DOCX
        else:
            raise ValueError(
                f"Unsupported file extension: {ext}. "
                f"Supported extensions are: pdf, txt, docx"
            )

    def _process_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process a PDF file and extract its content and metadata."""
        text = ""
        metadata = {"page_count": 0}
        
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            metadata["page_count"] = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

        metadata["word_count"] = len(text.split())
        return text, metadata

    def _process_txt(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process a text file and extract its content and metadata."""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        metadata = {
            "page_count": 1,
            "word_count": len(text.split())
        }
        return text, metadata

    def _process_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process a DOCX file and extract its content and metadata."""
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
                text += "\n"
        
        metadata = {
            "page_count": len(doc.sections),  # Approximate page count based on sections
            "word_count": len(text.split())
        }
        return text, metadata

    def _process_pdf_in_memory(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """Process a PDF file from memory and extract its content and metadata."""
        text = ""
        metadata = {"page_count": 0}
        
        # Create a file-like object from bytes
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        metadata["page_count"] = len(pdf_reader.pages)
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"

        metadata["word_count"] = len(text.split())
        return text, metadata

    def _process_txt_in_memory(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """Process a text file from memory and extract its content and metadata."""
        # Create a file-like object from bytes
        text_file = BytesIO(file_content)
        text = text_file.read().decode('utf-8')
        
        metadata = {
            "page_count": 1,
            "word_count": len(text.split())
        }
        return text, metadata

    def _process_docx_in_memory(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """Process a DOCX file from memory and extract its content and metadata."""
        # Create a file-like object from bytes
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
                text += "\n"
        
        metadata = {
            "page_count": len(doc.sections),  # Approximate page count based on sections
            "word_count": len(text.split())
        }
        return text, metadata

    def chunk_document(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[DocumentChunk]:
        """Split document text into overlapping chunks."""
        chunks = []
        words = text.split()
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += 1

            if current_size >= chunk_size:
                chunk_text = " ".join(current_chunk)
                chunks.append(DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    document_id="",  # Will be set by the caller
                    content=chunk_text,
                    metadata={"word_count": len(chunk_text.split())}
                ))
                
                # Keep overlap words for the next chunk
                current_chunk = current_chunk[-overlap:]
                current_size = len(current_chunk)

        # Add the last chunk if it exists
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(DocumentChunk(
                chunk_id=str(uuid.uuid4()),
                document_id="",  # Will be set by the caller
                content=chunk_text,
                metadata={"word_count": len(chunk_text.split())}
            ))

        return chunks 